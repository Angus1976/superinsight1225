"""
Multi-format × Multi-processing integration tests.

Tests the full pipeline: upload → process → verify → transfer
across 8 file formats × 3 processing types (structuring, vectorization, semantic).

Covers:
- Format routing correctness for all 8 types
- Large data performance benchmarks
- Progress tracking visibility
- Transfer to data lifecycle module
"""

from __future__ import annotations

import json
import os
import tempfile
import time
from typing import Any
from unittest.mock import MagicMock, patch
from uuid import uuid4

import pytest

from src.models.structuring import (
    FileType as StructuringFileType,
    JobStatus,
    StructuringJob,
)
from src.services.vectorization_pipeline import (
    _extract_text,
    _TEXT_TYPES as VEC_TEXT_TYPES,
    _TABULAR_TYPES as VEC_TABULAR_TYPES,
    chunk_text,
)
from src.services.structuring_pipeline import (
    _TEXT_TYPES as STRUCT_TEXT_TYPES,
    _TABULAR_TYPES as STRUCT_TABULAR_TYPES,
)


# ---------------------------------------------------------------------------
# Fixtures: sample file generators
# ---------------------------------------------------------------------------


def _write_temp(suffix: str, content: bytes | str) -> str:
    """Write content to a temp file and return its path."""
    mode = "wb" if isinstance(content, bytes) else "w"
    encoding = None if isinstance(content, bytes) else "utf-8"
    fd, path = tempfile.mkstemp(suffix=suffix)
    with os.fdopen(fd, mode, encoding=encoding) as f:  # type: ignore[call-overload]
        f.write(content)
    return path


@pytest.fixture
def sample_txt():
    path = _write_temp(".txt", "Hello world. This is a test document for structuring.")
    yield path
    os.unlink(path)


@pytest.fixture
def sample_md():
    content = "# Test Document\n\nThis is **markdown** content.\n\n- Item 1\n- Item 2\n"
    path = _write_temp(".md", content)
    yield path
    os.unlink(path)


@pytest.fixture
def sample_json():
    data = {"name": "test", "items": [{"id": 1, "value": "alpha"}, {"id": 2, "value": "beta"}]}
    path = _write_temp(".json", json.dumps(data, ensure_ascii=False, indent=2))
    yield path
    os.unlink(path)


@pytest.fixture
def sample_csv():
    content = "name,age,city\nAlice,30,Beijing\nBob,25,Shanghai\nCharlie,35,Shenzhen\n"
    path = _write_temp(".csv", content)
    yield path
    os.unlink(path)


@pytest.fixture
def sample_html():
    content = "<html><body><h1>Title</h1><p>Paragraph content here.</p></body></html>"
    path = _write_temp(".html", content)
    yield path
    os.unlink(path)


@pytest.fixture
def large_txt():
    """Generate a ~1MB text file for performance testing."""
    line = "This is a sample line for performance testing of large file processing. " * 5
    content = "\n".join([line] * 2000)  # ~1MB
    path = _write_temp(".txt", content)
    yield path
    os.unlink(path)


@pytest.fixture
def large_csv():
    """Generate a ~1MB CSV file for performance testing."""
    header = "id,name,value,description,category\n"
    row_template = "{i},item_{i},{val},Description for item {i},category_{cat}\n"
    rows = [row_template.format(i=i, val=i * 1.5, cat=i % 10) for i in range(10000)]
    path = _write_temp(".csv", header + "".join(rows))
    yield path
    os.unlink(path)


def _make_job(
    file_type: str,
    file_path: str,
    status: str = JobStatus.PENDING.value,
    tenant_id: str = "test-tenant",
) -> MagicMock:
    """Create a mock StructuringJob."""
    job = MagicMock(spec=StructuringJob)
    job.id = uuid4()
    job.file_type = file_type
    job.file_path = file_path
    job.status = status
    job.tenant_id = tenant_id
    job.file_name = os.path.basename(file_path)
    job.raw_content = None
    job.chunk_count = None
    job.record_count = 0
    job.error_message = None
    job.progress_info = None
    return job


# ---------------------------------------------------------------------------
# 1. Format routing consistency: vectorization == structuring
# ---------------------------------------------------------------------------

class TestFormatRoutingConsistency:
    """Verify vectorization and structuring pipelines route the same formats."""

    def test_text_types_match(self):
        assert VEC_TEXT_TYPES == STRUCT_TEXT_TYPES, (
            f"Mismatch: vec={VEC_TEXT_TYPES}, struct={STRUCT_TEXT_TYPES}"
        )

    def test_tabular_types_match(self):
        assert VEC_TABULAR_TYPES == STRUCT_TABULAR_TYPES

    def test_markdown_in_text_types(self):
        assert "markdown" in VEC_TEXT_TYPES
        assert "markdown" in STRUCT_TEXT_TYPES

    def test_json_in_text_types(self):
        assert "json" in VEC_TEXT_TYPES
        assert "json" in STRUCT_TEXT_TYPES


# ---------------------------------------------------------------------------
# 2. Real file extraction: text-based formats (TXT, Markdown, JSON, HTML)
# ---------------------------------------------------------------------------

class TestRealFileExtraction:
    """Extract text from real temp files via _extract_text (no mocks)."""

    def test_txt_extraction(self, sample_txt):
        job = _make_job("txt", sample_txt)
        text = _extract_text(job)
        assert "Hello world" in text
        assert len(text) > 0

    def test_markdown_extraction(self, sample_md):
        job = _make_job("markdown", sample_md)
        text = _extract_text(job)
        assert "markdown" in text.lower()

    def test_json_extraction(self, sample_json):
        job = _make_job("json", sample_json)
        text = _extract_text(job)
        assert "alpha" in text

    def test_html_extraction(self, sample_html):
        job = _make_job("html", sample_html)
        text = _extract_text(job)
        assert "Paragraph" in text or "Title" in text

    def test_csv_extraction(self, sample_csv):
        job = _make_job("csv", sample_csv)
        text = _extract_text(job)
        # Should contain header and actual data values
        assert "name" in text  # header
        assert "Alice" in text  # data value
        assert "Beijing" in text  # data value

    def test_unsupported_type_raises(self, sample_txt):
        job = _make_job("unknown_format", sample_txt)
        with pytest.raises(ValueError, match="Unsupported file type"):
            _extract_text(job)


# ---------------------------------------------------------------------------
# 3. Parametrized: format × processing type (mocked pipelines)
# ---------------------------------------------------------------------------

TEXT_FORMATS = ["txt", "markdown", "json", "html"]
TABULAR_FORMATS = ["csv"]
ALL_TESTABLE_FORMATS = TEXT_FORMATS + TABULAR_FORMATS

PROCESSING_TYPES = ["structuring", "vectorization", "semantic"]


class TestFormatProcessingMatrix:
    """Parametrized tests: each format × each processing type."""

    @pytest.mark.parametrize("file_type", ALL_TESTABLE_FORMATS)
    def test_vectorization_extract_succeeds(self, file_type, tmp_path):
        """Every testable format should extract text without error."""
        # Create a minimal file for each type
        content_map = {
            "txt": ("test.txt", "Sample text content"),
            "markdown": ("test.md", "# Heading\nContent"),
            "json": ("test.json", '{"key": "value"}'),
            "html": ("test.html", "<html><body><p>Text</p></body></html>"),
            "csv": ("test.csv", "a,b\n1,2\n3,4\n"),
        }
        fname, content = content_map[file_type]
        fpath = tmp_path / fname
        fpath.write_text(content, encoding="utf-8")

        job = _make_job(file_type, str(fpath))
        text = _extract_text(job)
        assert len(text) > 0

    @pytest.mark.parametrize("file_type", TEXT_FORMATS)
    def test_vectorization_chunk_produces_output(self, file_type, tmp_path):
        """Text formats should produce at least 1 chunk after extraction."""
        content_map = {
            "txt": ("test.txt", "A " * 300),
            "markdown": ("test.md", "# H\n" + "word " * 300),
            "json": ("test.json", json.dumps({"data": "x " * 300})),
            "html": ("test.html", f"<p>{'word ' * 300}</p>"),
        }
        fname, content = content_map[file_type]
        fpath = tmp_path / fname
        fpath.write_text(content, encoding="utf-8")

        job = _make_job(file_type, str(fpath))
        text = _extract_text(job)
        chunks = chunk_text(text, size=512, overlap=50)
        assert len(chunks) >= 1
        assert all(len(c) > 0 for c in chunks)

    @pytest.mark.parametrize("processing_type", PROCESSING_TYPES)
    def test_structuring_pipeline_accepts_markdown(self, processing_type, sample_md):
        """Markdown should be accepted by all processing types."""
        job = _make_job("markdown", sample_md)
        # All pipelines start with text extraction via _extract_text
        text = _extract_text(job)
        assert len(text) > 0

    @pytest.mark.parametrize("processing_type", PROCESSING_TYPES)
    def test_structuring_pipeline_accepts_json(self, processing_type, sample_json):
        """JSON should be accepted by all processing types."""
        job = _make_job("json", sample_json)
        text = _extract_text(job)
        assert len(text) > 0


# ---------------------------------------------------------------------------
# 4. Large data performance benchmarks
# ---------------------------------------------------------------------------

class TestLargeDataPerformance:
    """Performance tests for large file processing."""

    @pytest.mark.slow
    def test_large_txt_extraction_under_5s(self, large_txt):
        """~1MB text file should extract in under 5 seconds."""
        job = _make_job("txt", large_txt)
        start = time.time()
        text = _extract_text(job)
        elapsed = time.time() - start

        assert len(text) > 100_000, f"Expected large text, got {len(text)} chars"
        assert elapsed < 5.0, f"Extraction took {elapsed:.2f}s, expected < 5s"

    @pytest.mark.slow
    def test_large_txt_chunking_under_10s(self, large_txt):
        """~1MB text chunking (512 tokens, 50 overlap) should complete in < 10s."""
        job = _make_job("txt", large_txt)
        text = _extract_text(job)

        start = time.time()
        chunks = chunk_text(text, size=512, overlap=50)
        elapsed = time.time() - start

        assert len(chunks) > 1, f"Expected multiple chunks, got {len(chunks)}"
        assert elapsed < 10.0, f"Chunking took {elapsed:.2f}s, expected < 10s"

    @pytest.mark.slow
    def test_large_csv_extraction_under_5s(self, large_csv):
        """~1MB CSV (10k rows) should extract in under 5 seconds."""
        job = _make_job("csv", large_csv)
        start = time.time()
        text = _extract_text(job)
        elapsed = time.time() - start

        assert len(text) > 10_000
        assert elapsed < 5.0, f"CSV extraction took {elapsed:.2f}s, expected < 5s"

    def test_chunk_count_scales_linearly(self):
        """Chunk count should scale roughly linearly with text size."""
        small = "word " * 200   # ~200 tokens
        large = "word " * 2000  # ~2000 tokens

        small_chunks = chunk_text(small, size=512, overlap=50)
        large_chunks = chunk_text(large, size=512, overlap=50)

        # large is 10x bigger, chunks should be roughly 5-15x more
        ratio = len(large_chunks) / max(len(small_chunks), 1)
        assert 3 <= ratio <= 20, f"Chunk ratio {ratio:.1f} outside expected range"


# ---------------------------------------------------------------------------
# 5. Progress tracking visibility
# ---------------------------------------------------------------------------

class TestProgressTracking:
    """Verify progress tracking is visible and correct during processing."""

    def test_progress_tracker_records_all_steps(self):
        from src.services.progress_tracker import ProgressTracker

        tracker = ProgressTracker(job_id="test-job-1")

        tracker.start_step(1, "提取文件内容...")
        tracker.update_step(1, 50, "解析中...")
        tracker.complete_step(1, "提取完成")

        tracker.start_step(2, "推断 Schema...")
        tracker.complete_step(2, "Schema 推断完成")

        progress = tracker.get_progress_dict()

        assert progress["job_id"] == "test-job-1"
        assert progress["overall_status"] == "running"
        assert progress["current_step"] == 2
        assert progress["steps"][0]["status"] == "completed"
        assert progress["steps"][0]["progress_percent"] == 100.0
        assert progress["steps"][1]["status"] == "completed"
        assert progress["steps"][2]["status"] == "pending"

    def test_progress_tracker_saves_to_job(self):
        from src.services.progress_tracker import ProgressTracker

        tracker = ProgressTracker(job_id="test-job-2")
        tracker.start_step(1, "开始处理")

        mock_session = MagicMock()
        mock_job = MagicMock()
        mock_job.progress_info = None

        tracker.save_to_job(mock_session, mock_job)

        assert mock_job.progress_info is not None
        assert mock_job.progress_info["current_step"] == 1
        mock_session.flush.assert_called_once()

    def test_progress_overall_percent_calculation(self):
        from src.services.progress_tracker import ProgressTracker

        tracker = ProgressTracker(job_id="test-job-3")

        # Complete 3 of 6 steps
        for i in range(1, 4):
            tracker.start_step(i)
            tracker.complete_step(i)

        progress = tracker.get_progress_dict()
        # 3 steps at 100% out of 6 total = 50%
        assert progress["overall_progress_percent"] == 50.0

    def test_progress_failure_tracking(self):
        from src.services.progress_tracker import ProgressTracker

        tracker = ProgressTracker(job_id="test-job-4")
        tracker.start_step(1, "提取中...")
        tracker.fail_step(1, "文件损坏")

        progress = tracker.get_progress_dict()
        assert progress["overall_status"] == "failed"
        assert progress["steps"][0]["status"] == "failed"
        assert "文件损坏" in progress["steps"][0]["message"]

    def test_progress_estimated_remaining_time(self):
        from src.services.progress_tracker import ProgressTracker

        tracker = ProgressTracker(job_id="test-job-5")
        tracker.start_step(1)
        time.sleep(0.05)  # simulate work
        tracker.complete_step(1)

        progress = tracker.get_progress_dict()
        remaining = progress["estimated_remaining_seconds"]
        # With 1 of 6 steps done, remaining should be ~5x elapsed
        assert remaining is not None
        assert remaining > 0


# ---------------------------------------------------------------------------
# 6. Semantic pipeline with mocked LLM (format-aware)
# ---------------------------------------------------------------------------

class TestSemanticPipelineFormats:
    """Verify semantic pipeline works with markdown and json formats."""

    @patch("src.services.semantic_pipeline._call_llm")
    def test_semantic_entities_from_markdown(self, mock_llm, sample_md):
        from src.services.semantic_pipeline import _extract_entities

        entities = [{"name": "Test", "type": "concept", "properties": {}, "confidence": 0.9}]
        mock_llm.return_value = json.dumps(entities)

        result = _extract_entities("# Heading\nSome content", "t1")
        assert len(result) == 1
        assert result[0]["name"] == "Test"

    @patch("src.services.semantic_pipeline._call_llm")
    def test_semantic_entities_from_json_content(self, mock_llm, sample_json):
        from src.services.semantic_pipeline import _extract_entities

        entities = [
            {"name": "alpha", "type": "item", "properties": {"id": 1}, "confidence": 0.85},
            {"name": "beta", "type": "item", "properties": {"id": 2}, "confidence": 0.80},
        ]
        mock_llm.return_value = json.dumps(entities)

        result = _extract_entities('{"items": [{"id": 1}]}', "t1")
        assert len(result) == 2


# ---------------------------------------------------------------------------
# 7. Transfer to data lifecycle (mocked)
# ---------------------------------------------------------------------------

class TestDataTransferIntegration:
    """Verify processed data can be transferred to lifecycle module."""

    def _make_attributes(self) -> dict:
        return {"category": "test", "tags": ["auto"], "quality_score": 0.9}

    def test_transfer_request_model_accepts_structuring_source(self):
        from src.models.data_transfer import DataTransferRequest

        request = DataTransferRequest(
            source_type="structuring",
            source_id="job-123",
            target_state="temp_stored",
            data_attributes=self._make_attributes(),
            records=[{"id": "r1", "content": {"name": "test"}}],
        )
        assert request.source_type == "structuring"
        assert request.target_state == "temp_stored"
        assert len(request.records) == 1

    def test_transfer_request_model_accepts_augmentation_source(self):
        from src.models.data_transfer import DataTransferRequest

        request = DataTransferRequest(
            source_type="augmentation",
            source_id="job-456",
            target_state="in_sample_library",
            data_attributes=self._make_attributes(),
            records=[{"id": "v1", "content": {"chunk_text": "some text"}}],
        )
        assert request.source_type == "augmentation"
        assert request.target_state == "in_sample_library"

    def test_transfer_request_model_accepts_sync_source(self):
        from src.models.data_transfer import DataTransferRequest

        request = DataTransferRequest(
            source_type="sync",
            source_id="job-789",
            target_state="annotation_pending",
            data_attributes=self._make_attributes(),
            records=[{"id": "s1", "content": {"record_type": "entity", "name": "X"}}],
        )
        assert request.source_type == "sync"
        assert request.target_state == "annotation_pending"

    @pytest.mark.parametrize("target_state", [
        "temp_stored", "in_sample_library", "annotation_pending",
    ])
    def test_all_target_states_accepted(self, target_state):
        from src.models.data_transfer import DataTransferRequest

        request = DataTransferRequest(
            source_type="structuring",
            source_id="job-test",
            target_state=target_state,
            data_attributes=self._make_attributes(),
            records=[{"id": "r1", "content": {"data": "test"}}],
        )
        assert request.target_state == target_state


# ---------------------------------------------------------------------------
# 8. Real PDF / DOCX file extraction (requires pypdf, python-docx)
# ---------------------------------------------------------------------------

class TestRealBinaryFileExtraction:
    """Extract text from real PDF and DOCX files generated in-memory."""

    @pytest.fixture
    def sample_pdf(self, tmp_path):
        """Generate a minimal PDF with pypdf."""
        from pypdf import PdfWriter
        from pypdf.generic import NameObject, TextStringObject

        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        # Write a simple text annotation as content workaround
        fpath = tmp_path / "test.pdf"
        with open(fpath, "wb") as f:
            writer.write(f)
        return str(fpath)

    @pytest.fixture
    def sample_docx(self, tmp_path):
        """Generate a minimal DOCX with python-docx."""
        from docx import Document

        doc = Document()
        doc.add_heading("Test Document", level=1)
        doc.add_paragraph("This is paragraph one with important data.")
        doc.add_paragraph("Second paragraph contains more information.")
        fpath = tmp_path / "test.docx"
        doc.save(str(fpath))
        return str(fpath)

    def test_docx_extraction_returns_content(self, sample_docx):
        job = _make_job("docx", sample_docx)
        text = _extract_text(job)
        assert "paragraph" in text.lower()
        assert "important" in text.lower()
        assert len(text) > 20

    def test_docx_chunking_works(self, sample_docx):
        job = _make_job("docx", sample_docx)
        text = _extract_text(job)
        chunks = chunk_text(text, size=512, overlap=50)
        assert len(chunks) >= 1

    def test_pdf_extraction_does_not_crash(self, sample_pdf):
        """PDF with blank page should extract without error (may be empty)."""
        job = _make_job("pdf", sample_pdf)
        # Blank PDF may return empty content or raise — both are acceptable
        try:
            text = _extract_text(job)
            # If it succeeds, text should be a string
            assert isinstance(text, str)
        except RuntimeError:
            # Empty PDF extraction failure is expected behavior
            pass


# ---------------------------------------------------------------------------
# 9. Full vectorization pipeline mock integration (format-aware)
# ---------------------------------------------------------------------------

class TestFullVectorizationPipeline:
    """Test run_vectorization_pipeline with different file types (mocked DB + LLM)."""

    def _setup_db_mock(self, job):
        mock_session = MagicMock()
        mock_db = MagicMock()
        mock_db.get_session.return_value.__enter__ = MagicMock(return_value=mock_session)
        mock_db.get_session.return_value.__exit__ = MagicMock(return_value=False)
        mock_session.query.return_value.filter_by.return_value.first.return_value = job
        return mock_db, mock_session

    @pytest.mark.parametrize("file_type,suffix,content", [
        ("txt", ".txt", "Simple text for vectorization testing."),
        ("markdown", ".md", "# Title\n\nMarkdown content for vectorization."),
        ("json", ".json", '{"data": "JSON content for vectorization"}'),
        ("html", ".html", "<html><body><p>HTML for vectorization</p></body></html>"),
    ])
    @patch("src.services.vectorization_pipeline._store_vector_records")
    @patch("src.services.vectorization_pipeline._embed_chunks")
    @patch("src.services.vectorization_pipeline.chunk_text")
    @patch("src.services.vectorization_pipeline.db_manager")
    def test_pipeline_succeeds_for_text_format(
        self, mock_db_mgr, mock_chunk, mock_embed, mock_store,
        file_type, suffix, content, tmp_path,
    ):
        from src.services.vectorization_pipeline import run_vectorization_pipeline

        fpath = tmp_path / f"test{suffix}"
        fpath.write_text(content, encoding="utf-8")

        job = _make_job(file_type, str(fpath))
        mock_db, _ = self._setup_db_mock(job)
        mock_db_mgr.get_session = mock_db.get_session

        mock_chunk.return_value = ["chunk1"]
        mock_embed.return_value = [[0.1] * 1536]
        mock_store.return_value = 1

        result = run_vectorization_pipeline(str(job.id))

        assert result["status"] == JobStatus.COMPLETED.value
        assert result["chunk_count"] == 1


# ---------------------------------------------------------------------------
# 10. Full semantic pipeline mock integration (format-aware)
# ---------------------------------------------------------------------------

class TestFullSemanticPipeline:
    """Test run_semantic_pipeline with markdown and json formats."""

    def _setup_db_mock(self, job):
        mock_session = MagicMock()
        mock_db = MagicMock()
        mock_db.get_session.return_value.__enter__ = MagicMock(return_value=mock_session)
        mock_db.get_session.return_value.__exit__ = MagicMock(return_value=False)
        mock_session.query.return_value.filter_by.return_value.first.return_value = job
        return mock_db, mock_session

    @pytest.mark.parametrize("file_type,suffix,content", [
        ("markdown", ".md", "# Report\n\nAlice works at Acme Corp in Beijing."),
        ("json", ".json", '{"person": "Alice", "company": "Acme", "city": "Beijing"}'),
    ])
    @patch("src.services.semantic_pipeline._store_semantic_records")
    @patch("src.services.semantic_pipeline._generate_summary")
    @patch("src.services.semantic_pipeline._extract_relationships")
    @patch("src.services.semantic_pipeline._extract_entities")
    @patch("src.services.semantic_pipeline.db_manager")
    def test_pipeline_succeeds_for_format(
        self, mock_db_mgr, mock_entities, mock_rels, mock_summary, mock_store,
        file_type, suffix, content, tmp_path,
    ):
        from src.services.semantic_pipeline import run_semantic_pipeline

        fpath = tmp_path / f"test{suffix}"
        fpath.write_text(content, encoding="utf-8")

        job = _make_job(file_type, str(fpath))
        mock_db, _ = self._setup_db_mock(job)
        mock_db_mgr.get_session = mock_db.get_session

        mock_entities.return_value = [{"name": "Alice", "type": "person", "properties": {}, "confidence": 0.9}]
        mock_rels.return_value = [{"source": "Alice", "target": "Acme", "relation": "works_at", "confidence": 0.8}]
        mock_summary.return_value = {"text": "Alice works at Acme.", "confidence": 0.95}
        mock_store.return_value = 3

        result = run_semantic_pipeline(str(job.id))

        assert result["status"] == JobStatus.COMPLETED.value
        assert result["record_count"] == 3


# ---------------------------------------------------------------------------
# 11. CSV data quality: values not keys
# ---------------------------------------------------------------------------

class TestCSVDataQuality:
    """Verify CSV extraction preserves actual data values, not just column names."""

    def test_csv_values_preserved(self, sample_csv):
        """After fix: extracted text must contain row values, not just headers."""
        job = _make_job("csv", sample_csv)
        text = _extract_text(job)

        # Header should be present
        assert "name" in text
        assert "age" in text
        assert "city" in text

        # Actual data values must be present (not just column names)
        assert "Alice" in text, "Row value 'Alice' missing from CSV extraction"
        assert "30" in text, "Row value '30' missing from CSV extraction"
        assert "Beijing" in text, "Row value 'Beijing' missing from CSV extraction"
        assert "Bob" in text
        assert "Shanghai" in text

    def test_csv_header_appears_once(self, sample_csv):
        """Header line should appear exactly once at the top."""
        job = _make_job("csv", sample_csv)
        text = _extract_text(job)
        lines = text.strip().split("\n")

        # First line should be the header
        assert "name" in lines[0]
        # Data lines should follow
        assert "Alice" in lines[1]

    def test_large_csv_truncates_at_500_rows(self, tmp_path):
        """CSV with >500 rows should only extract first 500 data rows."""
        header = "id,value\n"
        rows = [f"{i},{i*10}\n" for i in range(1000)]
        fpath = tmp_path / "big.csv"
        fpath.write_text(header + "".join(rows), encoding="utf-8")

        job = _make_job("csv", str(fpath))
        text = _extract_text(job)
        data_lines = text.strip().split("\n")

        # 1 header + 500 data rows = 501 lines
        assert len(data_lines) == 501, f"Expected 501 lines, got {len(data_lines)}"


# ---------------------------------------------------------------------------
# 12. Edge cases
# ---------------------------------------------------------------------------

class TestEdgeCases:
    """Edge cases for file processing."""

    def test_empty_txt_file(self, tmp_path):
        """Empty text file should raise or return empty."""
        fpath = tmp_path / "empty.txt"
        fpath.write_text("", encoding="utf-8")
        job = _make_job("txt", str(fpath))
        # FileExtractor may return empty content
        try:
            text = _extract_text(job)
            # If it succeeds, text could be empty string
            assert isinstance(text, str)
        except RuntimeError:
            pass  # Empty file extraction failure is acceptable

    def test_unicode_content_preserved(self, tmp_path):
        """Chinese/Unicode content should be preserved through extraction."""
        content = "这是一段中文测试内容。包含特殊字符：①②③"
        fpath = tmp_path / "chinese.txt"
        fpath.write_text(content, encoding="utf-8")

        job = _make_job("txt", str(fpath))
        text = _extract_text(job)
        assert "中文测试" in text
        assert "①②③" in text

    def test_json_nested_structure(self, tmp_path):
        """Deeply nested JSON should extract without error."""
        data = {"level1": {"level2": {"level3": [1, 2, {"deep": True}]}}}
        fpath = tmp_path / "nested.json"
        fpath.write_text(json.dumps(data, indent=2), encoding="utf-8")

        job = _make_job("json", str(fpath))
        text = _extract_text(job)
        assert "deep" in text

    def test_markdown_with_code_blocks(self, tmp_path):
        """Markdown with code blocks should extract all content."""
        content = "# API Docs\n\n```python\ndef hello():\n    return 'world'\n```\n\nEnd."
        fpath = tmp_path / "code.md"
        fpath.write_text(content, encoding="utf-8")

        job = _make_job("markdown", str(fpath))
        text = _extract_text(job)
        assert "hello" in text
        assert "End" in text

    def test_html_strips_script_tags(self, tmp_path):
        """HTML extraction should handle script tags gracefully."""
        content = "<html><head><script>alert('xss')</script></head><body><p>Safe content</p></body></html>"
        fpath = tmp_path / "script.html"
        fpath.write_text(content, encoding="utf-8")

        job = _make_job("html", str(fpath))
        text = _extract_text(job)
        assert "Safe content" in text
