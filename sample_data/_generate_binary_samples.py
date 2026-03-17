"""生成连锁超市样本数据的二进制文件（PDF / DOCX / Excel）"""
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


# ── 1. Excel: 门店月度销售明细 ──────────────────────────────
def generate_excel():
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()

    # Sheet 1: 门店汇总
    ws1 = wb.active
    ws1.title = "门店月度汇总"
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="264653", end_color="264653", fill_type="solid")
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    headers1 = ["门店编号", "门店名称", "月销售额(万元)", "月交易笔数", "客单价(元)",
                 "会员占比", "毛利率", "同比增长", "环比增长", "员工数"]
    ws1.append(headers1)
    for cell in ws1[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
        cell.border = thin_border

    stores = [
        ["STORE-001", "朝阳路旗舰店", 456.82, 38580, 118.49, "69.4%", "24.3%", "+8.2%", "+3.1%", 45],
        ["STORE-002", "中关村店", 325.60, 28230, 115.24, "72.1%", "23.8%", "+5.1%", "+1.8%", 32],
        ["STORE-003", "望京店", 268.45, 24360, 110.17, "65.3%", "22.9%", "+3.7%", "+2.4%", 28],
        ["STORE-004", "通州万达店", 217.30, 20340, 106.71, "58.7%", "21.5%", "-1.2%", "-0.5%", 24],
        ["STORE-005", "大兴店", 170.25, 16020, 106.35, "61.2%", "22.1%", "+2.4%", "+1.2%", 20],
        ["STORE-006", "回龙观社区店", 116.80, 12690, 92.00, "74.8%", "25.6%", "+12.6%", "+5.3%", 15],
    ]
    for row in stores:
        ws1.append(row)
    for row in ws1.iter_rows(min_row=2, max_row=7):
        for cell in row:
            cell.border = thin_border
            cell.alignment = Alignment(horizontal="center")

    # Sheet 2: 品类分析
    ws2 = wb.create_sheet("品类销售分析")
    headers2 = ["品类", "销售额(万元)", "占比", "毛利率", "动销率", "库存周转天数", "TOP商品"]
    ws2.append(headers2)
    for cell in ws2[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    categories = [
        ["乳制品", 247.35, "15.9%", "22.3%", "94.2%", 8, "伊利纯牛奶"],
        ["粮油", 213.69, "13.7%", "18.5%", "88.7%", 25, "金龙鱼调和油"],
        ["饮料", 197.67, "12.7%", "35.2%", "91.5%", 12, "农夫山泉"],
        ["日用品", 175.02, "11.2%", "28.7%", "82.3%", 35, "维达抽纸"],
        ["休闲食品", 156.36, "10.1%", "32.1%", "86.9%", 18, "卫龙辣条"],
        ["方便食品", 124.89, "8.0%", "26.4%", "79.8%", 22, "康师傅牛肉面"],
        ["肉制品", 109.30, "7.0%", "20.1%", "90.1%", 10, "双汇火腿肠"],
        ["冷冻食品", 93.45, "6.0%", "24.8%", "76.5%", 30, "三全水饺"],
        ["酒水", 87.12, "5.6%", "30.5%", "68.2%", 45, "青岛啤酒"],
        ["调味品", 78.90, "5.1%", "35.8%", "85.4%", 40, "海天酱油"],
        ["其他", 71.47, "4.6%", "27.3%", "72.1%", 28, "-"],
    ]
    for row in categories:
        ws2.append(row)

    # Sheet 3: 每日流水（模拟30天）
    ws3 = wb.create_sheet("每日销售流水")
    headers3 = ["日期", "交易笔数", "销售额(元)", "退货额(元)", "净销售额(元)", "客单价(元)"]
    ws3.append(headers3)
    for cell in ws3[1]:
        cell.font = header_font
        cell.fill = header_fill

    import random
    random.seed(42)
    for day in range(1, 32):
        if day > 28 and day <= 31:
            date_str = f"2026-03-{day:02d}" if day <= 31 else None
            if day == 29 or day == 30 or day == 31:
                date_str = f"2026-03-{day:02d}"
        else:
            date_str = f"2026-03-{day:02d}"
        txn = random.randint(3800, 5200)
        revenue = round(txn * random.uniform(95, 125), 2)
        refund = round(revenue * random.uniform(0.005, 0.02), 2)
        net = round(revenue - refund, 2)
        avg = round(net / txn, 2)
        ws3.append([date_str, txn, revenue, refund, net, avg])

    # 调整列宽
    for ws in [ws1, ws2, ws3]:
        for col in ws.columns:
            max_len = max(len(str(cell.value or "")) for cell in col)
            ws.column_dimensions[col[0].column_letter].width = max(max_len + 4, 12)

    path = os.path.join(OUTPUT_DIR, "连锁超市_月度经营报表.xlsx")
    wb.save(path)
    print(f"✅ Excel: {path}")


# ── 2. DOCX: 员工培训手册 ──────────────────────────────────
def generate_docx():
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    # 标题
    title = doc.add_heading("万家乐连锁超市", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading("收银员操作手册（V3.2）", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("编制部门：运营管理中心\n生效日期：2026年3月1日\n密级：内部公开")
    doc.add_page_break()

    # 目录
    doc.add_heading("目录", level=1)
    toc_items = [
        "第一章 岗位职责与行为规范",
        "第二章 POS系统操作流程",
        "第三章 支付方式处理",
        "第四章 会员服务操作",
        "第五章 退换货处理流程",
        "第六章 异常情况应急处理",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()

    # 第一章
    doc.add_heading("第一章 岗位职责与行为规范", level=1)
    doc.add_heading("1.1 岗位职责", level=2)
    duties = [
        "准确、快速完成商品扫码与结算",
        "核验商品价格，发现异常及时报告值班主管",
        "妥善保管收银台备用金（标准：500元零钱）",
        "每日营业结束后完成交班对账",
        "保持收银台区域整洁，商品陈列规范",
        "主动为顾客提供购物袋、小票等服务",
    ]
    for d in duties:
        doc.add_paragraph(d, style="List Bullet")

    doc.add_heading("1.2 仪容仪表要求", level=2)
    doc.add_paragraph("统一着工装，佩戴工牌，保持整洁。女员工淡妆上岗，男员工不留长发。"
                       "禁止在收银台区域饮食、使用手机。")

    # 第二章
    doc.add_heading("第二章 POS系统操作流程", level=1)
    doc.add_heading("2.1 开机与登录", level=2)
    steps = [
        ("步骤1", "开启POS主机电源，等待系统启动（约45秒）"),
        ("步骤2", "输入员工编号（EMP-XXXX）和密码"),
        ("步骤3", "确认当前日期和班次（早班/晚班）"),
        ("步骤4", "清点备用金并在系统中确认金额"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "步骤"
    hdr[1].text = "操作说明"
    for step, desc in steps:
        row = table.add_row().cells
        row[0].text = step
        row[1].text = desc

    doc.add_heading("2.2 商品扫码", level=2)
    doc.add_paragraph(
        '将商品条形码对准扫码枪红光区域，听到"嘀"声表示扫码成功。'
        "如遇无法识别的条码，可手动输入13位条形码数字。"
        "散装商品需先在电子秤称重，打印价签后扫码录入。"
    )

    doc.add_heading("2.3 常用快捷键", level=2)
    keys_table = doc.add_table(rows=1, cols=3)
    keys_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr2 = keys_table.rows[0].cells
    hdr2[0].text = "快捷键"
    hdr2[1].text = "功能"
    hdr2[2].text = "使用场景"
    shortcuts = [
        ("F1", "帮助", "查看操作指南"),
        ("F2", "挂单", "顾客临时离开，暂存当前订单"),
        ("F3", "取单", "恢复挂起的订单"),
        ("F5", "会员查询", "输入手机号查询会员信息"),
        ("F8", "折扣", "输入主管授权的折扣码"),
        ("F10", "结算", "进入支付选择界面"),
        ("F12", "交班", "当班结束，打印交班报表"),
    ]
    for key, func, scene in shortcuts:
        row = keys_table.add_row().cells
        row[0].text = key
        row[1].text = func
        row[2].text = scene

    # 第三章
    doc.add_heading("第三章 支付方式处理", level=1)
    payments = {
        "现金支付": "收取现金 → 系统输入实收金额 → 确认找零 → 将找零和小票一并交给顾客",
        "微信/支付宝": "选择扫码支付 → 顾客出示付款码 → 扫码枪扫描 → 等待支付成功提示音",
        "银行卡": "选择银行卡支付 → 顾客插卡/挥卡 → 输入密码 → 等待POS机打印凭条",
        "会员积分抵扣": "F5查询会员 → 确认可用积分 → 输入抵扣积分数（100积分=1元）→ 差额用其他方式支付",
    }
    for method, flow in payments.items():
        doc.add_heading(method, level=2)
        doc.add_paragraph(flow)

    # 第四章
    doc.add_heading("第四章 会员服务操作", level=1)
    doc.add_paragraph(
        "会员等级体系：\n"
        "• 普通卡：累计消费 0-2,000元\n"
        "• 银卡：累计消费 2,001-10,000元，享98折\n"
        "• 金卡：累计消费 10,001-50,000元，享95折，积分1.5倍\n"
        "• 钻石卡：累计消费 50,001元以上，享88折，积分2倍，专属客服"
    )

    # 第五章
    doc.add_heading("第五章 退换货处理流程", level=1)
    doc.add_paragraph(
        "1. 顾客出示购物小票（7天内有效）\n"
        "2. 检查商品完整性，确认符合退换条件\n"
        "3. 在POS系统选择「退货」功能，扫描小票条码\n"
        "4. 选择退货商品，输入退货原因\n"
        "5. 值班主管刷卡授权（退货金额＞50元需主管授权）\n"
        "6. 原路退款，打印退货凭证\n"
        "7. 退货商品交由理货员处理"
    )

    # 第六章
    doc.add_heading("第六章 异常情况应急处理", level=1)
    emergencies = [
        ("POS死机", "长按电源键10秒强制关机，等待30秒后重启。如仍无法恢复，联系IT部（分机8888）"),
        ("停电", "立即停止收银，引导顾客有序等待。使用UPS备用电源保存当前数据"),
        ("假币识别", "使用验钞机复核，确认假币后礼貌告知顾客，必要时通知安保"),
        ("顾客投诉", "保持冷静，认真倾听。无法现场解决的，引导至服务台或呼叫值班主管"),
        ("系统价格错误", "暂停该商品销售，通知信息部核实。已售出的按较低价格执行"),
    ]
    for title_text, handling in emergencies:
        p = doc.add_paragraph()
        run = p.add_run(f"【{title_text}】")
        run.bold = True
        p.add_run(f"\n{handling}")

    path = os.path.join(OUTPUT_DIR, "连锁超市_收银员操作手册.docx")
    doc.save(path)
    print(f"✅ DOCX: {path}")


# ── 3. PDF: 食品安全检查报告 ──────────────────────────────
def generate_pdf():
    from fpdf import FPDF
    import os

    # 查找可用的中文字体
    font_path = None
    candidates = [
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
        "/System/Library/Fonts/STHeiti Medium.ttc",
    ]
    for fp in candidates:
        if os.path.exists(fp):
            font_path = fp
            break

    if not font_path:
        print("⚠️  未找到中文字体，PDF 将使用 ASCII 内容生成")
        _generate_pdf_ascii()
        return

    pdf = FPDF()
    pdf.add_page()
    pdf.add_font("Chinese", "", font_path, uni=True)
    pdf.set_font("Chinese", size=18)
    pdf.cell(0, 15, "万家乐连锁超市", ln=True, align="C")
    pdf.set_font("Chinese", size=14)
    pdf.cell(0, 12, "食品安全巡检报告", ln=True, align="C")
    pdf.ln(5)

    pdf.set_font("Chinese", size=10)
    info_lines = [
        "报告编号：FS-2026-0316-001",
        "巡检日期：2026年3月16日",
        "巡检门店：朝阳路旗舰店（STORE-001）",
        "巡检人员：刘卫生（食品安全主管）",
        "陪同人员：陈店长",
    ]
    for line in info_lines:
        pdf.cell(0, 7, line, ln=True)

    pdf.ln(5)
    pdf.set_font("Chinese", size=12)
    pdf.cell(0, 10, "一、巡检项目及结果", ln=True)
    pdf.ln(2)

    pdf.set_font("Chinese", size=9)
    col_widths = [15, 50, 55, 30, 40]
    headers = ["序号", "检查项目", "检查标准", "结果", "备注"]
    for i, h in enumerate(headers):
        pdf.cell(col_widths[i], 8, h, border=1, align="C")
    pdf.ln()

    items = [
        ["1", "冷藏柜温度", "0~4°C", "合格(2.8°C)", ""],
        ["2", "冷冻柜温度", "-18°C以下", "合格(-20°C)", ""],
        ["3", "熟食操作间温度", "≤25°C", "合格(23°C)", ""],
        ["4", "员工健康证", "在有效期内", "合格", "2人下月到期"],
        ["5", "食品标签标识", "清晰完整", "合格", ""],
        ["6", "临期食品管理", "专区陈列并标识", "合格", "已设专柜"],
        ["7", "消毒记录", "每日记录完整", "合格", ""],
        ["8", "防鼠防虫设施", "完好无破损", "整改中", "后门挡鼠板缺失"],
        ["9", "进货台账", "索证索票齐全", "合格", ""],
        ["10", "废弃物处理", "分类存放日产日清", "合格", ""],
        ["11", "食品留样", "125g以上保存48h", "合格", ""],
        ["12", "操作人员着装", "工帽口罩手套齐全", "合格", ""],
    ]
    for row in items:
        for i, val in enumerate(row):
            pdf.cell(col_widths[i], 7, val, border=1, align="C")
        pdf.ln()

    pdf.ln(5)
    pdf.set_font("Chinese", size=12)
    pdf.cell(0, 10, "二、问题及整改要求", ln=True)
    pdf.set_font("Chinese", size=10)
    issues = [
        "1. 后门挡鼠板缺失：要求3月18日前完成更换，由工程部负责。",
        "2. 2名员工健康证将于4月到期：提醒相关员工4月1日前完成体检续办。",
        "3. 望京店冷藏柜温度波动问题已上报设备部，建议48小时内完成检修。",
    ]
    for issue in issues:
        pdf.cell(0, 7, issue, ln=True)

    pdf.ln(5)
    pdf.set_font("Chinese", size=12)
    pdf.cell(0, 10, "三、总体评价", ln=True)
    pdf.set_font("Chinese", size=10)
    pdf.multi_cell(0, 7,
        "本次巡检共检查12个大项，合格11项，整改中1项，合格率91.7%。"
        "整体食品安全管理状况良好，门店卫生环境达标。"
        "需重点关注防鼠设施维护和员工健康证到期续办工作。"
        "建议下次巡检重点复查本次整改项目。"
    )

    pdf.ln(8)
    pdf.set_font("Chinese", size=10)
    pdf.cell(0, 7, "巡检人签字：________________    日期：2026年3月16日", ln=True)
    pdf.cell(0, 7, "店长签字：  ________________    日期：2026年3月16日", ln=True)

    path = os.path.join(OUTPUT_DIR, "连锁超市_食品安全巡检报告.pdf")
    pdf.output(path)
    print(f"✅ PDF: {path}")


def _generate_pdf_ascii():
    """Fallback: generate English PDF if no Chinese font available."""
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=16)
    pdf.cell(0, 12, "WanJiaLe Supermarket Chain", ln=True, align="C")
    pdf.set_font("Helvetica", size=13)
    pdf.cell(0, 10, "Food Safety Inspection Report", ln=True, align="C")
    pdf.ln(5)
    pdf.set_font("Helvetica", size=10)
    pdf.cell(0, 7, "Report ID: FS-2026-0316-001 | Date: 2026-03-16", ln=True)
    pdf.cell(0, 7, "Store: Chaoyang Flagship (STORE-001)", ln=True)
    pdf.cell(0, 7, "Inspector: Liu Weisheng (Food Safety Manager)", ln=True)
    pdf.ln(5)
    pdf.set_font("Helvetica", size=9)
    col_w = [10, 45, 50, 30, 55]
    for i, h in enumerate(["No.", "Item", "Standard", "Result", "Remark"]):
        pdf.cell(col_w[i], 8, h, border=1, align="C")
    pdf.ln()
    rows = [
        ["1", "Fridge Temp", "0~4C", "Pass(2.8C)", ""],
        ["2", "Freezer Temp", "<-18C", "Pass(-20C)", ""],
        ["3", "Deli Room Temp", "<=25C", "Pass(23C)", ""],
        ["4", "Health Certs", "Valid", "Pass", "2 expiring next month"],
        ["5", "Food Labels", "Clear & complete", "Pass", ""],
        ["6", "Near-expiry Mgmt", "Separate display", "Pass", ""],
        ["7", "Sanitize Records", "Daily complete", "Pass", ""],
        ["8", "Pest Control", "Intact barriers", "Fixing", "Back door plate missing"],
        ["9", "Purchase Ledger", "Docs complete", "Pass", ""],
        ["10", "Waste Disposal", "Sorted & daily", "Pass", ""],
    ]
    for row in rows:
        for i, v in enumerate(row):
            pdf.cell(col_w[i], 7, v, border=1, align="C")
        pdf.ln()
    path = os.path.join(OUTPUT_DIR, "连锁超市_食品安全巡检报告.pdf")
    pdf.output(path)
    print(f"✅ PDF (ASCII fallback): {path}")


if __name__ == "__main__":
    generate_excel()
    generate_docx()
    generate_pdf()
    print("\n🎉 所有二进制样本文件生成完毕！")
